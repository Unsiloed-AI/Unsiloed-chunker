#!/usr/bin/env python3
"""
Unsiloed Chunker Test Suite - Clean & Simple

Tests all chunking strategies with multiple document formats.
Downloads sample documents automatically if none provided.

Usage:
    python3 test_chunker.py                    # Auto-download and test
    python3 test_chunker.py document.pdf       # Test specific file
"""

import os
import sys
import time
import json
import requests
from pathlib import Path

# Add current directory to path for imports
sys.path.insert(0, str(Path(__file__).parent))

def create_sample_files():
    """Create sample files for testing."""
    print("📝 Creating sample files...")
    
    # Sample text content
    sample_text = """
# Sample Document for Testing

## Introduction
This is a sample document for testing the Unsiloed chunker with various document formats.

## Section 1: Document Types Supported
The chunker now supports multiple document formats:
- PDF documents (.pdf)
- Word documents (.docx) 
- PowerPoint presentations (.pptx)
- Excel spreadsheets (.xlsx)
- Plain text files (.txt)

## Section 2: Chunking Strategies
Each format can be processed with different chunking strategies:

### Fixed Chunking
Splits text into fixed-size chunks with optional overlap.
Best for: Uniform processing, simple text analysis.

### Paragraph Chunking  
Splits text at paragraph boundaries (double newlines).
Best for: Maintaining paragraph structure, document sections.

### Heading Chunking
Splits text at heading markers (# headers, numbered sections, etc).
Best for: Document structure preservation, chapter-based processing.

### Page Chunking (PDF only)
Extracts each PDF page as a separate chunk.
Best for: Page-level analysis, document scanning workflows.

### Semantic Chunking (Advanced)
Uses AI to understand document structure and create meaningful chunks.
Best for: High-quality RAG applications, semantic search.

## Conclusion
This multi-format support makes the chunker versatile for various document processing workflows.
"""
    
    created_files = []
    
    # Create text file
    text_file = Path("sample_test.txt")
    with open(text_file, 'w', encoding='utf-8') as f:
        f.write(sample_text.strip())
    print(f"✅ Created: {text_file}")
    created_files.append(text_file)
    
    # Try to download a simple PDF
    try:
        print("📥 Downloading sample PDF...")
        url = "https://www.w3.org/WAI/ER/tests/xhtml/testfiles/resources/pdf/dummy.pdf"
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        
        pdf_file = Path("sample_test.pdf")
        with open(pdf_file, 'wb') as f:
            f.write(response.content)
        print(f"✅ Downloaded: {pdf_file} ({len(response.content):,} bytes)")
        created_files.append(pdf_file)
        
    except Exception as e:
        print(f"⚠️ Could not download PDF: {e}")
    
    # Create DOCX if library available
    try:
        import docx
        print("📄 Creating sample DOCX...")
        doc = docx.Document()
        doc.add_heading('Sample DOCX Document', 0)
        doc.add_paragraph('This is a sample Word document for testing chunking strategies.')
        doc.add_heading('Features', level=1)
        doc.add_paragraph('The chunker can extract text while maintaining document structure.')
        
        docx_file = Path("sample_test.docx")
        doc.save(str(docx_file))
        print(f"✅ Created: {docx_file}")
        created_files.append(docx_file)
        
    except ImportError:
        print("ℹ️ python-docx not available (pip install python-docx)")
    
    return created_files

def extract_text_from_file(file_path):
    """Extract text from various document formats."""
    suffix = file_path.suffix.lower()
    
    try:
        if suffix == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
                
        elif suffix == '.pdf':
            import PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
                
        elif suffix == '.docx':
            import docx
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
            
        else:
            print(f"⚠️ Unsupported file format: {suffix}")
            return None
            
    except Exception as e:
        print(f"❌ Error extracting text from {file_path}: {e}")
        return None

def test_fixed_chunking(text, file_path):
    """Test fixed-size chunking."""
    print("\n🧩 Testing Fixed-Size Chunking...")
    
    try:
        from Unsiloed.utils.chunking import fixed_size_chunking
        
        start_time = time.time()
        chunks = fixed_size_chunking(text, chunk_size=500, overlap=50)
        duration = time.time() - start_time
        
        print(f"  ✅ Success: {len(chunks)} chunks in {duration:.3f}s")
        print(f"  📊 Avg chunk size: {sum(len(c['text']) for c in chunks) / len(chunks):.0f} chars")
        
        return {'status': 'success', 'chunks': len(chunks), 'time': duration}
        
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return {'status': 'error', 'error': str(e)}

def test_paragraph_chunking(text, file_path):
    """Test paragraph chunking."""
    print("\n📄 Testing Paragraph Chunking...")
    
    try:
        from Unsiloed.utils.chunking import paragraph_chunking
        
        start_time = time.time()
        chunks = paragraph_chunking(text)
        duration = time.time() - start_time
        
        print(f"  ✅ Success: {len(chunks)} chunks in {duration:.3f}s")
        
        return {'status': 'success', 'chunks': len(chunks), 'time': duration}
        
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return {'status': 'error', 'error': str(e)}

def test_heading_chunking(text, file_path):
    """Test heading chunking."""
    print("\n📋 Testing Heading Chunking...")
    
    try:
        from Unsiloed.utils.chunking import heading_chunking
        
        start_time = time.time()
        chunks = heading_chunking(text)
        duration = time.time() - start_time
        
        print(f"  ✅ Success: {len(chunks)} chunks in {duration:.3f}s")
        
        return {'status': 'success', 'chunks': len(chunks), 'time': duration}
        
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return {'status': 'error', 'error': str(e)}

def test_page_chunking(file_path):
    """Test page chunking (PDF only)."""
    if not str(file_path).lower().endswith('.pdf'):
        print("\n📖 Skipping Page Chunking (not a PDF)")
        return {'status': 'skipped', 'reason': 'Not a PDF file'}
    
    print("\n📖 Testing Page Chunking...")
    
    try:
        from Unsiloed.utils.chunking import page_based_chunking
        
        start_time = time.time()
        chunks = page_based_chunking(str(file_path))
        duration = time.time() - start_time
        
        print(f"  ✅ Success: {len(chunks)} pages in {duration:.3f}s")
        
        return {'status': 'success', 'chunks': len(chunks), 'time': duration}
        
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return {'status': 'error', 'error': str(e)}

def test_semantic_chunking(file_path):
    """Test semantic chunking."""
    print("\n🧠 Testing Semantic Chunking...")
    
    has_openai = bool(os.getenv('OPENAI_API_KEY'))
    print(f"  🔑 OpenAI API Key: {'Available' if has_openai else 'Not found (will use fallback)'}")
    
    try:
        from Unsiloed.utils.chunking import semantic_chunking
        
        start_time = time.time()
        result = semantic_chunking(str(file_path), max_concurrent_calls=3)
        duration = time.time() - start_time
        
        # Handle both dict and list returns
        if isinstance(result, dict):
            chunks = result.get('chunks', [])
            pages = len(result.get('image_dimensions', []))
        else:
            chunks = result
            pages = 1
        
        print(f"  ✅ Success: {len(chunks)} chunks in {duration:.3f}s")
        if pages > 0:
            print(f"  ⏱️ Performance: {duration/pages:.3f}s per page")
        
        return {'status': 'success', 'chunks': len(chunks), 'time': duration, 'pages': pages}
        
    except Exception as e:
        print(f"  ❌ Error: {e}")
        return {'status': 'error', 'error': str(e)}

def run_tests(file_path):
    """Run all tests on a file."""
    print(f"\n{'='*60}")
    print(f"🎯 Testing: {file_path.name}")
    print(f"📏 Size: {file_path.stat().st_size:,} bytes")
    print(f"📄 Format: {file_path.suffix.upper()}")
    print(f"{'='*60}")
    
    results = {'file': str(file_path), 'format': file_path.suffix.lower(), 'strategies': {}}
    
    # Extract text content for text-based strategies
    print(f"📖 Extracting text from {file_path.suffix.upper()} file...")
    text_content = extract_text_from_file(file_path)
    
    if text_content:
        word_count = len(text_content.split())
        char_count = len(text_content)
        print(f"  ✅ Extracted: {char_count:,} characters, {word_count:,} words")
        
        # Run text-based chunking strategies
        results['strategies']['fixed'] = test_fixed_chunking(text_content, file_path)
        results['strategies']['paragraph'] = test_paragraph_chunking(text_content, file_path)
        results['strategies']['heading'] = test_heading_chunking(text_content, file_path)
    else:
        print(f"  ❌ Could not extract text from {file_path.suffix.upper()} file")
        # Skip text-based strategies
        for strategy in ['fixed', 'paragraph', 'heading']:
            results['strategies'][strategy] = {
                'status': 'skipped', 
                'reason': f'Could not extract text from {file_path.suffix.upper()}'
            }
    
    # Run file-based strategies
    results['strategies']['page'] = test_page_chunking(file_path)
    results['strategies']['semantic'] = test_semantic_chunking(file_path)
    
    return results

def print_summary(all_results):
    """Print test summary."""
    print(f"\n{'='*60}")
    print("📊 TEST SUMMARY")
    print(f"{'='*60}")
    
    total_time = sum(
        result.get('time', 0) 
        for file_results in all_results 
        for result in file_results['strategies'].values()
        if isinstance(result, dict) and 'time' in result
    )
    
    print(f"⏱️ Total time: {total_time:.2f}s")
    print(f"📁 Files tested: {len(all_results)}")
    
    # Strategy summary
    strategy_stats = {}
    for file_results in all_results:
        for strategy_name, strategy_result in file_results['strategies'].items():
            if strategy_name not in strategy_stats:
                strategy_stats[strategy_name] = {'success': 0, 'error': 0, 'skipped': 0}
            
            status = strategy_result.get('status', 'unknown')
            strategy_stats[strategy_name][status] = strategy_stats[strategy_name].get(status, 0) + 1
    
    print(f"\n📈 Strategy Results:")
    for strategy, stats in strategy_stats.items():
        total = sum(stats.values())
        success_rate = stats.get('success', 0) / total * 100 if total > 0 else 0
        print(f"  {strategy.upper():<12}: {stats.get('success', 0)}/{total} success ({success_rate:.0f}%)")
    
    # Save results
    results_file = Path("test_results.json")
    with open(results_file, 'w', encoding='utf-8') as f:
        json.dump(all_results, f, indent=2, default=str)
    
    print(f"\n💾 Results saved to: {results_file}")

def main():
    """Main function."""
    print("🚀 Unsiloed Chunker Test Suite")
    print("=" * 40)
    
    # Get file to test
    if len(sys.argv) > 1:
        file_path = Path(sys.argv[1])
        if not file_path.exists():
            print(f"❌ File not found: {file_path}")
            print(f"💡 Available files:")
            for item in Path.cwd().iterdir():
                if item.is_file() and item.suffix.lower() in ['.pdf', '.txt', '.docx', '.pptx', '.xlsx']:
                    print(f"    📄 {item.name}")
            print(f"\n💡 Or run without arguments to create sample files:")
            print(f"    python3 test_chunker.py")
            return
            
        test_files = [file_path]
    else:
        test_files = create_sample_files()
        if not test_files:
            print("❌ Could not create test files")
            return
    
    print(f"\n🔧 Available strategies:")
    print(f"  ✅ Fixed chunking      - No API needed")
    print(f"  ✅ Paragraph chunking  - No API needed") 
    print(f"  ✅ Heading chunking    - No API needed")
    print(f"  ✅ Page chunking       - No API needed (PDF only)")
    print(f"  🔑 Semantic chunking   - OpenAI API optional")
    
    # Run tests
    all_results = []
    try:
        for file_path in test_files:
            if file_path.exists():
                result = run_tests(file_path)
                all_results.append(result)
        
        if all_results:
            print_summary(all_results)
            print(f"\n🎉 Testing completed successfully!")
        else:
            print(f"❌ No files were tested")
        
    except KeyboardInterrupt:
        print(f"\n⚠️ Testing interrupted")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()