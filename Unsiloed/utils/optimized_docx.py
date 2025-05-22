"""
Optimized DOCX processing module with memory-efficient operations.

This module provides advanced capabilities for handling DOCX files:
1. Streaming extraction for large documents
2. Efficient text and image extraction
3. Advanced structure preservation
4. Memory optimized processing
"""
import os
import io
import gc
import re
import logging
import tempfile
from typing import Iterator, Dict, Union, BinaryIO, List, Any, Optional, Tuple, Generator
import concurrent.futures
from pathlib import Path

from Unsiloed.utils.document_cache import document_cache
from Unsiloed.utils.memory_profiling import MemoryProfiler
from Unsiloed.utils.exceptions import (
    DocxExtractionError, 
    DependencyError,
    UnsupportedOperationError
)

logger = logging.getLogger(__name__)

# Try to import specialized document libraries
# Import required document processing libraries
try:
    import docx
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logger.warning("python-docx not available, will use alternative extraction methods")

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False
    logger.warning("mammoth not available, will use alternative extraction methods")

# Import BeautifulSoup for HTML parsing
try:
    from bs4 import BeautifulSoup
    BEAUTIFULSOUP_AVAILABLE = True
except ImportError:
    BEAUTIFULSOUP_AVAILABLE = False
    logger.warning("BeautifulSoup not available, HTML structure parsing will be limited")


def get_optimal_docx_extractor(file_path: Union[str, BinaryIO], 
                               force_method: Optional[str] = None) -> str:
    """
    Determine the optimal DOCX extraction method based on file characteristics
    and available libraries.
    
    Args:
        file_path: Path to the DOCX file or file-like object
        force_method: Force a specific extraction method
        
    Returns:
        Best method to use for extraction ("docx", "mammoth")
        
    Raises:
        DependencyError: If no suitable DOCX extraction methods are available
        FileNotFoundError: If the specified file path does not exist
        UnsupportedOperationError: If the forced method is not available
    """
    # Check if file exists if it's a string path
    if isinstance(file_path, str) and not isinstance(file_path, io.BytesIO):
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
    
    # If method is forced, use it if available
    if force_method:
        if force_method == "docx":
            if not PYTHON_DOCX_AVAILABLE:
                raise DependencyError(
                    "python-docx library is required but not installed", 
                    dependency="python-docx"
                )
            return "docx"
        elif force_method == "mammoth":
            if not MAMMOTH_AVAILABLE:
                raise DependencyError(
                    "mammoth library is required but not installed", 
                    dependency="mammoth"
                )
            return "mammoth"
        else:
            raise UnsupportedOperationError(f"Unsupported DOCX extraction method: {force_method}")
    
    # Otherwise, determine the best method
    # For most purposes, mammoth produces better output with structure preserved
    if MAMMOTH_AVAILABLE:
        return "mammoth"
    elif PYTHON_DOCX_AVAILABLE:
        return "docx"
    else:
        raise DependencyError(
            "No DOCX extraction methods available. Please install either mammoth or python-docx",
            dependency="mammoth or python-docx"
        )


def extract_text_streaming_docx(
    file_path: Union[str, BinaryIO], 
    extraction_method: Optional[str] = None,
    batch_size: int = 100  # Increased batch size for better performance
) -> Generator[Dict[str, Union[str, int]], None, None]:
    """
    Stream text extraction from DOCX with optimized memory usage.
    
    Args:
        file_path: Path to the DOCX file or file-like object
        extraction_method: Force a specific extraction method
        batch_size: Number of paragraphs to process in each batch
        
    Yields:
        Dictionary with paragraph index and extracted text
        
    Raises:
        DocxExtractionError: If there's an error during DOCX extraction
        DependencyError: If required dependencies are missing
        FileNotFoundError: If the specified file path does not exist
    """
    # Generate a file identifier for logging
    file_id = os.path.basename(str(file_path)) if isinstance(file_path, str) else "bytesio_document"
    
    try:
        # Determine best extraction method
        method = extraction_method or get_optimal_docx_extractor(file_path)
        
        # Use a memory profiler to track usage during streaming extraction
        profiler = MemoryProfiler(f"docx_stream_extract_{file_id}")
        profiler.start()
        
        try:
            if method == "docx" and PYTHON_DOCX_AVAILABLE:
                try:
                    # Handle file-like object
                    doc = docx.Document(file_path)
                    
                    paragraphs = doc.paragraphs
                    total_paragraphs = len(paragraphs)
                    
                    # Process paragraphs in larger batches for better performance
                    # Use StringIO for more efficient string operations
                    current_batch = io.StringIO()
                    current_size = 0
                    batch_para_start = 0
                    batch_size_limit = 8192  # Larger batch size limit for better performance
                    
                    for para_num, para in enumerate(paragraphs):
                        text = para.text.strip()
                        if not text:
                            continue
                            
                        # Append to StringIO buffer for better performance
                        current_batch.write(text)
                        current_batch.write("\n")
                        current_size += len(text) + 1  # +1 for the newline
                        
                        if current_size >= batch_size_limit or para_num == total_paragraphs - 1:
                            text_content = current_batch.getvalue()
                            yield {
                                "paragraph": batch_para_start,
                                "text": text_content,
                                "paragraph_count": para_num - batch_para_start + 1
                            }
                            
                            # Reset batch
                            current_batch = io.StringIO()
                            current_size = 0
                            batch_para_start = para_num + 1
                            
                            # Explicit garbage collection only after larger batches
                            gc.collect()
                except Exception as e:
                    logger.error(f"Error processing DOCX with python-docx: {str(e)}")
                    raise DocxExtractionError(
                        f"Failed to extract text using python-docx: {str(e)}",
                        extraction_method="python-docx",
                        file_path=str(file_path) if isinstance(file_path, str) else None,
                        details={"error": str(e)}
                    ) from e
                    
            elif method == "mammoth" and MAMMOTH_AVAILABLE:
                # Handle file-like objects by writing to temporary file if needed
                temp_path = None
                
                try:
                    if isinstance(file_path, io.BytesIO):
                        # Create temporary file
                        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                            temp_path = temp_file.name
                            file_path.seek(0)
                            temp_file.write(file_path.read())
                        
                        with open(temp_path, "rb") as doc_file:
                            # Convert to HTML first to preserve structure
                            result = mammoth.convert_to_html(doc_file)
                            html = result.value
                    else:
                        # Direct file path processing
                        with open(file_path, "rb") as doc_file:
                            # Convert to HTML first to preserve structure
                            result = mammoth.convert_to_html(doc_file)
                            html = result.value
                    
                    # Parse HTML to extract text in batches
                    if not BEAUTIFULSOUP_AVAILABLE:
                        logger.warning("BeautifulSoup not available, falling back to basic text processing")
                        # Simple fallback when BeautifulSoup is not available
                        # Strip HTML tags crudely
                        text = html.replace("<", " ").replace(">", " ")
                        # Split by newlines or obvious paragraph boundaries to create batches
                        paragraphs = [p for p in re.split(r'</?p>|</?h[1-6]>|\n\n|\r\n\r\n', text) if p.strip()]
                        
                        # Process in batches
                        current_batch = io.StringIO()
                        current_size = 0
                        batch_start = 0
                        
                        for i, para in enumerate(paragraphs):
                            text = para.strip()
                            if not text:
                                continue
                                
                            current_batch.write(text)
                            current_batch.write("\n")
                            current_size += len(text) + 1
                            
                            if current_size >= batch_size_limit or i == len(paragraphs) - 1:
                                text_content = current_batch.getvalue()
                                yield {
                                    "paragraph": batch_start,
                                    "text": text_content,
                                    "paragraph_count": i - batch_start + 1
                                }
                                
                                # Reset batch
                                current_batch = io.StringIO()
                                current_size = 0
                                batch_start = i + 1
                                
                                # Garbage collect
                                gc.collect()
                    else:
                        # Use BeautifulSoup for better HTML parsing
                        soup = BeautifulSoup(html, 'html.parser')
                        elements = soup.find_all(['p', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'li'])
                        
                        # Process elements in larger batches for better performance
                        total_elements = len(elements)
                        batch_size_limit = 8192  # Larger batch size for better performance
                        
                        # Use StringIO for more efficient string operations
                        current_batch = io.StringIO()
                        current_size = 0
                        batch_elem_start = 0
                        
                        for elem_num, elem in enumerate(elements):
                            text = elem.get_text().strip()
                            if not text:
                                continue
                                
                            # Append to StringIO buffer for better performance
                            current_batch.write(text)
                            current_batch.write("\n")
                            current_size += len(text) + 1  # +1 for the newline
                            
                            if current_size >= batch_size_limit or elem_num == total_elements - 1:
                                text_content = current_batch.getvalue()
                                yield {
                                    "paragraph": batch_elem_start,
                                    "text": text_content,
                                    "element_count": elem_num - batch_elem_start + 1
                                }
                                
                                # Reset batch
                                current_batch = io.StringIO()
                                current_size = 0
                                batch_elem_start = elem_num + 1
                                
                                # Explicit garbage collection only after larger batches
                                gc.collect()
                except Exception as e:
                    logger.error(f"Error processing DOCX with mammoth: {str(e)}")
                    raise DocxExtractionError(
                        f"Failed to extract text using mammoth: {str(e)}",
                        extraction_method="mammoth",
                        file_path=str(file_path) if isinstance(file_path, str) else None,
                        details={"error": str(e)}
                    ) from e
                finally:
                    # Clean up temporary file
                    if temp_path:
                        try:
                            os.unlink(temp_path)
                        except Exception as e:
                            logger.warning(f"Failed to remove temporary DOCX file: {str(e)}")
            else:
                raise DependencyError(
                    "No suitable DOCX extraction method available",
                    dependency="python-docx or mammoth" 
                )
        finally:
            # Always stop profiling and log memory usage
            memory_stats = profiler.stop()
            logger.debug(f"Memory usage during DOCX extraction: {memory_stats}")
    except (DocxExtractionError, DependencyError) as e:
        # Re-raise these specific exceptions
        raise
    except Exception as e:
        # Catch any other exceptions and wrap them
        logger.error(f"Unexpected error during DOCX extraction: {str(e)}")
        raise DocxExtractionError(
            f"Unexpected error during DOCX extraction: {str(e)}",
            extraction_method=extraction_method,
            file_path=str(file_path) if isinstance(file_path, str) else None,
            details={"error": str(e), "error_type": str(type(e).__name__)}
        ) from e


def extract_text_docx_with_structure(
    file_path: Union[str, BinaryIO],
    extraction_method: Optional[str] = None
) -> Dict[str, Any]:
    """
    Extract text from DOCX with structure preservation and memory optimization.
    
    Args:
        file_path: Path to the DOCX file or file-like object
        extraction_method: Force a specific extraction method
        
    Returns:
        Dictionary with extracted structured content
    """
    # Check cache first
    cached_result = document_cache.get(file_path, "docx_structured")
    if cached_result:
        return cached_result
    
    # Determine best extraction method
    method = extraction_method or get_optimal_docx_extractor(file_path)
    
    # Profile memory usage
    with MemoryProfiler(f"docx_structured_{os.path.basename(str(file_path))}") as profiler:
        result = {
            "text": "",
            "paragraphs": [],
            "headings": [],
            "tables": [],
            "images": []
        }
        
        if method == "mammoth" and MAMMOTH_AVAILABLE:
            # Handle file-like objects by writing to temporary file if needed
            if isinstance(file_path, io.BytesIO):
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as temp_file:
                    temp_path = temp_file.name
                    file_path.seek(0)
                    temp_file.write(file_path.read())
                
                try:
                    with open(temp_path, "rb") as doc_file:
                        # Convert to HTML to preserve structure
                        conversion = mammoth.convert_to_html(doc_file)
                        html = conversion.value
                        
                        # Store warnings
                        result["warnings"] = [str(msg) for msg in conversion.messages]
                finally:
                    # Clean up temporary file
                    try:
                        os.unlink(temp_path)
                    except:
                        pass
            else:
                with open(file_path, "rb") as doc_file:
                    conversion = mammoth.convert_to_html(doc_file)
                    html = conversion.value
                    result["warnings"] = [str(msg) for msg in conversion.messages]
            
            # Parse HTML structure
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, 'html.parser')
            
            # Extract all text
            result["text"] = soup.get_text()
            
            # Extract paragraphs
            for i, p in enumerate(soup.find_all('p')):
                text = p.get_text().strip()
                if text:
                    result["paragraphs"].append({
                        "index": i,
                        "text": text
                    })
            
            # Extract headings
            for i, h in enumerate(soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])):
                text = h.get_text().strip()
                if text:
                    result["headings"].append({
                        "index": i,
                        "level": int(h.name[1]),
                        "text": text
                    })
            
            # Extract tables
            for i, table in enumerate(soup.find_all('table')):
                table_data = []
                for row in table.find_all('tr'):
                    row_data = []
                    for cell in row.find_all(['td', 'th']):
                        row_data.append(cell.get_text().strip())
                    table_data.append(row_data)
                
                if table_data:
                    result["tables"].append({
                        "index": i,
                        "data": table_data
                    })
        
        elif method == "docx" and PYTHON_DOCX_AVAILABLE:
            # Process with python-docx
            if isinstance(file_path, io.BytesIO):
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(file_path)
            
            full_text = []
            
            # Process paragraphs
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    result["paragraphs"].append({
                        "index": i,
                        "text": text,
                        "style": para.style.name if para.style else "Normal"
                    })
                    full_text.append(text)
                    
                    # Check if paragraph is a heading
                    if para.style and para.style.name.startswith("Heading"):
                        try:
                            level = int(para.style.name.replace("Heading ", ""))
                            result["headings"].append({
                                "index": i,
                                "level": level,
                                "text": text
                            })
                        except:
                            pass
            
            # Process tables in batches to manage memory
            for i, table in enumerate(doc.tables):
                table_data = []
                
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                
                result["tables"].append({
                    "index": i,
                    "data": table_data
                })
                
                # Add table text to full text
                for row in table_data:
                    full_text.append(" | ".join(row))
                
                # Free memory after processing each table
                if i % 5 == 0:  # Process tables in batches of 5
                    gc.collect()
            
            result["text"] = "\n\n".join(full_text)
            
        else:
            raise RuntimeError("No suitable DOCX extraction method available")
    
    # Store in cache for future use
    document_cache.set(file_path, "docx_structured", result)
    
    return result


def extract_docx_with_images(file_path: Union[str, BinaryIO]) -> Dict[str, Any]:
    """
    Extract both text and images from a DOCX document with memory optimization.
    
    Args:
        file_path: Path to the DOCX file or file-like object
        
    Returns:
        Dictionary with extracted text and image data
    """
    # Check if python-docx is available for image extraction
    if not PYTHON_DOCX_AVAILABLE:
        raise RuntimeError("python-docx is required for image extraction")
    
    # Check cache first
    cached_result = document_cache.get(file_path, "docx_with_images")
    if cached_result:
        return cached_result
    
    # Profile memory usage
    with MemoryProfiler(f"docx_images_{os.path.basename(str(file_path))}") as profiler:
        result = {
            "text": "",
            "paragraphs": [],
            "images": []
        }
        
        # Create temporary directory for images
        with tempfile.TemporaryDirectory() as temp_dir:
            # Load document
            if isinstance(file_path, io.BytesIO):
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(file_path)
            
            # Extract paragraphs
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
                    result["paragraphs"].append({
                        "index": i,
                        "text": text
                    })
            
            # Combine all text
            result["text"] = "\n\n".join(paragraphs)
            
            # Extract images - this is a bit complex with python-docx
            try:
                # Access internal relationships and extract images
                rels = doc.part.rels
                
                for rel in rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_part = rel.target_part
                            image_bytes = image_part.blob
                            content_type = image_part.content_type
                            
                            # Save image to temporary file to get dimensions
                            img_temp_path = os.path.join(temp_dir, f"img_{len(result['images'])}.{content_type.split('/')[-1]}")
                            
                            with open(img_temp_path, "wb") as img_file:
                                img_file.write(image_bytes)
                            
                            # Get image dimensions if PIL is available
                            width = 0
                            height = 0
                            try:
                                from PIL import Image
                                with Image.open(img_temp_path) as img:
                                    width, height = img.size
                            except ImportError:
                                pass
                            
                            # Add image info to results
                            result["images"].append({
                                "index": len(result["images"]),
                                "content_type": content_type,
                                "width": width,
                                "height": height,
                                "size": len(image_bytes)
                                # We don't include binary data to avoid bloating memory
                            })
                        except Exception as e:
                            logger.warning(f"Error extracting image from DOCX: {str(e)}")
            except Exception as e:
                logger.warning(f"Error accessing DOCX relationships: {str(e)}")
    
    # Store in cache for future use
    document_cache.set(file_path, "docx_with_images", result)
    
    return result
